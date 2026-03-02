"""Script per generare il documento di esempio del processo aziendale."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_sample_process():
    doc = Document()

    # Title
    title = doc.add_heading("Procedura di Richiesta Rimborso Spese", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Questa procedura descrive i passi necessari per richiedere il rimborso "
        "di spese aziendali sostenute dal dipendente. Seguire attentamente ogni "
        "passo e raccogliere tutta la documentazione richiesta."
    )

    # --- Step 1 ---
    doc.add_heading("Passo 1 — Dati del Dipendente", level=1)
    doc.add_paragraph(
        "Il dipendente deve fornire le seguenti informazioni personali:"
    )
    items = [
        "Nome e cognome completo",
        "Matricola dipendente (si trova nella busta paga, in alto a sinistra, "
        "codice di 6 cifre che inizia con 'M')",
        "Reparto / Dipartimento di appartenenza",
        "Nome del responsabile diretto (il proprio manager)",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    # --- Step 2 ---
    doc.add_heading("Passo 2 — Dettagli della Spesa", level=1)
    doc.add_paragraph(
        "Per ogni spesa da rimborsare, fornire i seguenti dettagli:"
    )
    items = [
        "Data della spesa (giorno/mese/anno)",
        "Importo totale in euro (IVA inclusa)",
        "Categoria della spesa: Viaggio, Vitto, Alloggio, Materiale, Formazione, Altro",
        "Descrizione breve della spesa (es. 'Pranzo con cliente XYZ a Milano')",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    # --- Step 3 ---
    doc.add_heading("Passo 3 — Documentazione di Supporto", level=1)
    doc.add_paragraph(
        "La documentazione è obbligatoria per procedere con il rimborso. "
        "Il dipendente deve confermare di avere:"
    )
    items = [
        "Scontrino fiscale o fattura originale (fotografare e conservare l'originale)",
        "Se la spesa supera i 250€: autorizzazione preventiva del responsabile "
        "(email o approvazione nel sistema gestionale, sezione 'Autorizzazioni Spese')",
        "Per le spese di viaggio: biglietto o ricevuta del mezzo di trasporto",
        "Per le spese di alloggio: conferma di prenotazione e ricevuta dell'hotel",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "NOTA: Se il dipendente non ha la documentazione, deve contattare "
        "il fornitore per richiedere una copia. Senza documentazione il rimborso "
        "non può essere elaborato."
    )

    # --- Step 4 ---
    doc.add_heading("Passo 4 — Codice Centro di Costo", level=1)
    doc.add_paragraph(
        "Ogni spesa deve essere associata a un centro di costo. Il dipendente deve fornire:"
    )
    items = [
        "Codice centro di costo (formato: CC-XXXX, dove XXXX è un numero a 4 cifre. "
        "Si trova nel portale HR aziendale, sezione 'Il mio profilo' > 'Dati contabili', "
        "oppure chiedere al proprio responsabile)",
        "Se la spesa è legata a un progetto specifico: codice progetto (formato: PRJ-XXXX, "
        "reperibile nel sistema gestionale sezione 'Progetti attivi')",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    # --- Step 5 ---
    doc.add_heading("Passo 5 — Modalità di Rimborso", level=1)
    doc.add_paragraph(
        "Il dipendente deve indicare la modalità preferita di rimborso:"
    )
    items = [
        "Accredito su busta paga (opzione predefinita, entro il mese successivo)",
        "Bonifico su conto corrente (solo se diverso da quello indicato per lo stipendio, "
        "in tal caso fornire IBAN completo)",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    # --- Step 6 ---
    doc.add_heading("Passo 6 — Riepilogo e Conferma", level=1)
    doc.add_paragraph(
        "Al termine della raccolta dati, l'agente deve presentare un riepilogo "
        "completo di tutte le informazioni raccolte. Il dipendente deve confermare "
        "che i dati sono corretti."
    )
    doc.add_paragraph(
        "Dopo la conferma, il dipendente dovrà:"
    )
    items = [
        "Accedere al portale HR (hr.azienda.it)",
        "Navigare a 'Rimborsi' > 'Nuova richiesta'",
        "Inserire i dati raccolti nel modulo online",
        "Allegare la documentazione di supporto",
        "Inviare la richiesta — riceverà una email di conferma con numero pratica",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "Tempi di elaborazione: 5-10 giorni lavorativi. Per verificare lo stato "
        "della richiesta, accedere al portale HR sezione 'Le mie richieste'."
    )

    doc.save("sample_process.docx")
    print("✅ Documento sample_process.docx creato con successo!")


if __name__ == "__main__":
    create_sample_process()
